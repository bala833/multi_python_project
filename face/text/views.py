from django.shortcuts import render,redirect
from django.http import HttpResponse,HttpResponseRedirect
from .models import ImageToText
from text.forms import ImageToTextForm
from rest_framework.response import Response
from face2.models import Face
from django.urls import reverse

from docx import Document
from PIL import Image
from pytesseract import image_to_string
import pytesseract
import PIL  
from curses import ascii
from PIL import Image
from django.shortcuts import get_list_or_404, get_object_or_404
from django.views.generic import TemplateView
from playsound import playsound
from django.contrib.auth.models import User

all_face = Face.objects.all().count()


def clean(text):
    return str(''.join(ascii.isprint(c) and c or ' ' for c in text)) 

pytesseract.pytesseract.tesseract_cmd = r"C:\\Program Files\\Tesseract-OCR\\tesseract.exe"

def imagetotext(request,template_name="image_to_text.html"):
  
    count_all_face = Face.objects.all().count()
    if request.method == 'POST': 
        form = ImageToTextForm(request.POST, request.FILES) 
        print(form)
        # image_ = request.POST.get('image',None)

        # print(image_,"eeeeeeeeeeeeeeeeeeeeeee")
        # if image_:
        if form.is_valid(): 
            img = form.cleaned_data.get('image')
            image = Image.open(img, mode='r')
            img_text = pytesseract.image_to_string(image)


            img_text = clean(img_text)
            print(img_text)
            form.save() 
            print((form.instance.pk),"asdfasdfasdfsdf")
            data = ImageToText.objects.get(id = form.instance.pk)
            # print(data,"eeeeeeeeeeeeeeeeee")
            data.text = img_text
            data.save()
            return redirect('success') 
    else: 
        form = ImageToTextForm() 
    return render(request, template_name, {'form' : form}) 


  
  
def success(request,template_name="list_page_img_to_txt.html"):
    count_all_face = Face.objects.all().count()
    img_txt_ = 'img_to_text'
    img_txt = ImageToText.objects.all()

    return render(request, template_name, locals()) 







def delete_text(request,id):
    # obj = get_object_or_404(ImageToText, pk=id)
    print(id,"dd")
    obj = ImageToText.objects.get(pk=id)
    if obj:
        obj.delete()
    # data.delete()

    return redirect('success')  

    # image = Image.open('/home/xarxa-15/Downloads/notes.jpeg', mode='r')
    # print(image)
    # bala = pytesseract.image_to_string(image)
    # print(bala)
    # # write in csv bala
    # document = Document()
    # document.add_heading('convet image to string', 0)
    # # cleaned_string = ''.join(c for c in bala if valid_xml_char_ordinal(c))
    # # print(cleaned_string,"kkkkkkkkkkkkkkkkkkk")
    # def clean(text):
    #     return str(''.join(
    #             ascii.isprint(c) and c or '?' for c in text
    #             )) 
    # bala = clean(bala)
    # p = document.add_paragraph(bala)
    # p.add_run('bold').bold = True
    # p.add_run(' and some ')
    # p.add_run('italic.').italic = True
    # document.add_page_break()
    # document.save('/home/xarxa-15/demo.docx')




def typing_test(request, template_name='time_speed.html'):
    time = 'time'

    return render(request, template_name, locals())



def play_song(request):
    play = playsound('/path/to/a/sound/file/you/want/to/play.mp3')
    return play

from django.contrib.auth import authenticate, login
from django.contrib.auth import logout

from django.contrib.auth import login as auth_login
def login_(request):
    count_all_face = Face.objects.all().count()
    user_name = request.POST.get('uname')
    user_pass = request.POST.get('psw')
    user = authenticate(username=user_name, password=user_pass)
    print(user,"invalid dataaaaaaaaa")
    if user is not None:
        login(request, user)
        if user.is_superuser:
            # print(request.user.is_authenticated)
            # print(request.user)
            # print(user.is_superuser,'is_superuser')
            success_url = reverse('success_login')
            return HttpResponseRedirect(success_url)
        else:
            print('nonnnnnnnnnnn')
    else:
        validation_error = 'Please Enter correct username and password'
        return render(request, locals())


    # if user.is_superuser:
    #     print(user.password)
    #     success_url = reverse('success_login')
    #     print(success_url,"jjjjjjjjjjjjjjj")
    #     return HttpResponseRedirect(success_url)
    #     # A backend authenticated the credentials
    # else:
    #     print('nonnnnnnn')
    # No backend authenticated the credentials

    # user = User.objects.filter(username=user_name,password=user_pass)
    # print(user,"kkkkkkkkkkkkkkkkkkkkk")
    # if user.is_superuser:
    #     print('super user')
    # else:
    #     print('user is not found')
    return render(request, locals())



from django.contrib.auth import get_user_model
def success_login(request, template_name="super_admin_dashboad.html"):
    count_all_face = Face.objects.all().count()
    bala = 'login'
    User = get_user_model()
    print(User.username)
    print(request.user.get_username())
    print(request.user.username,"kkkkkk")
    print(request.user.is_authenticated, request.user.is_superuser)

    return render(request,template_name, locals())





def logout_view(request):
    count_all_face = Face.objects.all().count()
    logout(request)
    success_url = reverse('home')
    return HttpResponseRedirect(success_url)
    # Redirect to a success page.